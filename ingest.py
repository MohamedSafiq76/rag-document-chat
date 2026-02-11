"""
Document Ingestion Module
Handles parsing of PDF, DOCX, and CSV files into chunked LangChain Documents.
"""

import os
import pandas as pd
from docx import Document as DocxDocument
from pypdf import PdfReader
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter


# ---------------------------------------------------------------------------
# Text Splitter — shared across all file types
# ---------------------------------------------------------------------------
text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=500,
    chunk_overlap=100,
    separators=["\n\n", "\n", ". ", " ", ""],
)


# ---------------------------------------------------------------------------
# Individual Parsers
# ---------------------------------------------------------------------------

def _load_pdf(file) -> list[Document]:
    """Parse a PDF file and return chunked Documents with page metadata."""
    reader = PdfReader(file)
    documents = []
    for page_num, page in enumerate(reader.pages, start=1):
        text = page.extract_text()
        if text and text.strip():
            chunks = text_splitter.split_text(text)
            for idx, chunk in enumerate(chunks):
                documents.append(Document(
                    page_content=chunk,
                    metadata={
                        "source": file.name,
                        "page": page_num,
                        "chunk": idx + 1,
                        "type": "pdf",
                    },
                ))
    return documents


def _load_docx(file) -> list[Document]:
    """Parse a DOCX file and return chunked Documents."""
    doc = DocxDocument(file)
    full_text = "\n".join(para.text for para in doc.paragraphs if para.text.strip())
    if not full_text.strip():
        return []

    chunks = text_splitter.split_text(full_text)
    documents = []
    for idx, chunk in enumerate(chunks):
        documents.append(Document(
            page_content=chunk,
            metadata={
                "source": file.name,
                "chunk": idx + 1,
                "type": "docx",
            },
        ))
    return documents


def _load_csv(file) -> list[Document]:
    """Parse a CSV file and return chunked Documents (one per row group)."""
    df = pd.read_csv(file)
    # Convert each row to a readable string representation
    rows_text = []
    columns = df.columns.tolist()
    for row_idx, row in df.iterrows():
        row_str = " | ".join(f"{col}: {val}" for col, val in zip(columns, row.values))
        rows_text.append(row_str)

    # Join rows and then chunk
    full_text = "\n".join(rows_text)
    chunks = text_splitter.split_text(full_text)

    documents = []
    for idx, chunk in enumerate(chunks):
        documents.append(Document(
            page_content=chunk,
            metadata={
                "source": file.name,
                "chunk": idx + 1,
                "type": "csv",
                "columns": ", ".join(columns),
            },
        ))
    return documents


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".csv"}

_LOADER_MAP = {
    ".pdf": _load_pdf,
    ".docx": _load_docx,
    ".csv": _load_csv,
}


def ingest_file(file) -> list[Document]:
    """
    Ingest a single uploaded file and return a list of chunked Documents.

    Parameters
    ----------
    file : streamlit.UploadedFile
        The file object from Streamlit's file_uploader.

    Returns
    -------
    list[Document]
        Chunked documents ready for embedding.
    """
    ext = os.path.splitext(file.name)[1].lower()
    loader = _LOADER_MAP.get(ext)
    if loader is None:
        raise ValueError(f"Unsupported file type: {ext}. Supported: {SUPPORTED_EXTENSIONS}")
    return loader(file)


def ingest_files(files) -> list[Document]:
    """
    Ingest multiple uploaded files and return all chunked Documents.

    Parameters
    ----------
    files : list[streamlit.UploadedFile]
        List of uploaded file objects.

    Returns
    -------
    list[Document]
        Combined chunked documents from all files.
    """
    all_docs = []
    for file in files:
        all_docs.extend(ingest_file(file))
    return all_docs
